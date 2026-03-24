import os
from PyPDF2 import PdfReader
from docx import Document


class FileExtractor:
    """Extracts text from various file types with improved content handling"""

    SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".md", ".csv", ".rtf"}
    MAX_CONTENT_SIZE = 15000  # chars per file

    @staticmethod
    def extract_text(file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return FileExtractor._read_pdf(file_path)
        elif ext in (".doc", ".docx"):
            return FileExtractor._read_docx(file_path)
        elif ext in (".txt", ".md", ".csv", ".rtf"):
            return FileExtractor._read_text(file_path)
        return ""

    @staticmethod
    def _read_pdf(path: str) -> str:
        try:
            reader = PdfReader(path)
            pages = []
            total_chars = 0
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if total_chars + len(page_text) > FileExtractor.MAX_CONTENT_SIZE:
                    remaining = FileExtractor.MAX_CONTENT_SIZE - total_chars
                    if remaining > 0:
                        pages.append(f"[Page {i + 1}]\n{page_text[:remaining]}")
                    pages.append(f"\n... (truncated, {len(reader.pages) - i - 1} more pages)")
                    break
                pages.append(f"[Page {i + 1}]\n{page_text}")
                total_chars += len(page_text)
            return "\n\n".join(pages)
        except Exception as e:
            return f"[PDF extraction error: {str(e)}]"

    @staticmethod
    def _read_docx(path: str) -> str:
        try:
            doc = Document(path)
            parts = []
            total_chars = 0

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if total_chars + len(text) > FileExtractor.MAX_CONTENT_SIZE:
                    parts.append("... (truncated)")
                    break
                parts.append(text)
                total_chars += len(text)

            # Extract tables if present
            for table in doc.tables[:5]:
                table_text = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(cells))
                if table_text:
                    table_str = "\n".join(table_text)
                    if total_chars + len(table_str) <= FileExtractor.MAX_CONTENT_SIZE:
                        parts.append(f"\n[Table]\n{table_str}")
                        total_chars += len(table_str)

            return "\n".join(parts)
        except Exception as e:
            return f"[DOCX extraction error: {str(e)}]"

    @staticmethod
    def _read_text(path: str) -> str:
        try:
            with open(path, 'r', errors='ignore') as f:
                content = f.read(FileExtractor.MAX_CONTENT_SIZE)
            if os.path.getsize(path) > FileExtractor.MAX_CONTENT_SIZE:
                content += "\n... (truncated)"
            return content
        except Exception as e:
            return f"[Text extraction error: {str(e)}]"
